#!/usr/bin/env python3
"""
Script to convert VPS_DEPLOYMENT_GUIDE.md to DOCX format
Requires: pip install python-docx markdown beautifulsoup4
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import re
import os

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style the hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    
    run.append(rPr)
    run_text = OxmlElement('w:t')
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)

    paragraph._p.append(hyperlink)
    return hyperlink

def convert_md_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to DOCX."""
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html = markdown.markdown(md_content, extensions=['codehilite', 'fenced_code', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add custom styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    
    # Heading styles
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(12)
    
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(16)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(14)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(6)
    
    # Code style
    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)
    
    # Process content
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'code', 'ul', 'ol', 'hr']):
        if element.name == 'h1':
            if 'HƯỚNG DẪN DEPLOY' in element.get_text():
                # Main title
                title = doc.add_paragraph(element.get_text(), style='CustomTitle')
            else:
                p = doc.add_paragraph(element.get_text(), style='CustomH1')
                
        elif element.name == 'h2':
            p = doc.add_paragraph(element.get_text(), style='CustomH2')
            
        elif element.name == 'h3':
            p = doc.add_paragraph(element.get_text(), style='CustomH3')
            
        elif element.name == 'h4':
            p = doc.add_paragraph(element.get_text(), style='CustomH3')
            
        elif element.name == 'pre':
            # Code blocks
            code_text = element.get_text()
            for line in code_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line, style='CustomCode')
                    
        elif element.name == 'p':
            # Regular paragraphs
            p = doc.add_paragraph()
            for child in element.children:
                if hasattr(child, 'name'):
                    if child.name == 'code':
                        # Inline code
                        run = p.add_run(child.get_text())
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    elif child.name == 'strong' or child.name == 'b':
                        # Bold text
                        run = p.add_run(child.get_text())
                        run.font.bold = True
                    elif child.name == 'em' or child.name == 'i':
                        # Italic text
                        run = p.add_run(child.get_text())
                        run.font.italic = True
                    elif child.name == 'a':
                        # Links
                        href = child.get('href', '')
                        text = child.get_text()
                        if href.startswith('http'):
                            add_hyperlink(p, href, text)
                        else:
                            p.add_run(text)
                    else:
                        p.add_run(child.get_text())
                else:
                    # Text nodes
                    text = str(child).strip()
                    if text:
                        p.add_run(text)
                        
        elif element.name in ['ul', 'ol']:
            # Lists
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text(), style='List Bullet' if element.name == 'ul' else 'List Number')
                
        elif element.name == 'hr':
            # Horizontal rule
            doc.add_paragraph('─' * 50)
    
    # Add page break before sections
    section_titles = ['CHUẨN BỊ VPS', 'CÀI ĐẶT MÔI TRƯỜNG', 'DEPLOY ỨNG DỤNG', 
                     'CẤU HÌNH NGINX', 'SSL CERTIFICATE', 'DATABASE SETUP',
                     'MONITORING & LOGS', 'BACKUP & MAINTENANCE']
    
    for paragraph in doc.paragraphs:
        if any(title in paragraph.text for title in section_titles):
            paragraph._element.addprevious(paragraph._element.getparent().add_page_break())
            break
    
    # Save document
    doc.save(docx_file_path)
    print(f"✅ Successfully converted {md_file_path} to {docx_file_path}")

def main():
    """Main function."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'VPS_DEPLOYMENT_GUIDE.md')
    docx_file = os.path.join(script_dir, 'VPS_DEPLOYMENT_GUIDE.docx')
    
    if not os.path.exists(md_file):
        print(f"❌ File không tồn tại: {md_file}")
        return
    
    try:
        convert_md_to_docx(md_file, docx_file)
        print(f"🎉 Đã tạo file DOCX: {docx_file}")
    except Exception as e:
        print(f"❌ Lỗi khi convert: {str(e)}")

if __name__ == "__main__":
    main() 